import os
import tempfile
from pathlib import Path
from typing import Dict, Any, List
import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
# weasyprint rimosso per problemi di dipendenze
from ebooklib import epub
from bs4 import BeautifulSoup

from .models import DSAProfile, ExportResult

logger = logging.getLogger(__name__)

class ExportManager:
    def __init__(self):
        self.fonts_path = Path(__file__).parent.parent.parent / "assets" / "fonts"
        self.templates_path = Path(__file__).parent.parent.parent / "assets" / "templates"
    
    async def export_document(
        self,
        structured_content: Dict[str, Any],
        dsa_profile: DSAProfile,
        format_type: str,
        output_directory: str,
        original_filename: str
    ) -> str:
        """Esporta il documento nel formato specificato"""
        try:
            # Crea la directory di output se non esiste
            os.makedirs(output_directory, exist_ok=True)
            
            # Genera il nome del file di output
            base_name = Path(original_filename).stem
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            if format_type == 'docx':
                return await self._export_docx(
                    structured_content, dsa_profile, output_directory, base_name, timestamp
                )
            elif format_type == 'pdf':
                return await self._export_pdf(
                    structured_content, dsa_profile, output_directory, base_name, timestamp
                )
            elif format_type == 'epub':
                return await self._export_epub(
                    structured_content, dsa_profile, output_directory, base_name, timestamp
                )
            else:
                raise ValueError(f"Formato non supportato: {format_type}")
                
        except Exception as e:
            logger.error(f"Errore nell'export {format_type}: {e}")
            raise
    
    async def _export_docx(
        self,
        structured_content: Dict[str, Any],
        dsa_profile: DSAProfile,
        output_directory: str,
        base_name: str,
        timestamp: str
    ) -> str:
        """Esporta in formato DOCX"""
        try:
            # Crea un nuovo documento
            doc = Document()
            
            # Configura gli stili
            self._setup_docx_styles(doc, dsa_profile)
            
            # Aggiungi il titolo
            if structured_content.get('title'):
                title = doc.add_heading(structured_content['title'], level=1)
                title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Aggiungi le sezioni
            for section in structured_content.get('sections', []):
                # Aggiungi il titolo della sezione
                heading = doc.add_heading(section['title'], level=min(section['level'], 6))
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Aggiungi il contenuto della sezione
                for content_line in section['content']:
                    if content_line.strip():
                        para = doc.add_paragraph(content_line)
                        para.style = 'DSA Body'
            
            # Aggiungi i paragrafi
            for paragraph in structured_content.get('paragraphs', []):
                if paragraph.strip():
                    para = doc.add_paragraph(paragraph)
                    para.style = 'DSA Body'
            
            # Aggiungi gli elenchi
            for list_item in structured_content.get('lists', []):
                if list_item['items']:
                    if list_item['type'] == 'bullet':
                        for item in list_item['items']:
                            para = doc.add_paragraph(item['text'], style='List Bullet')
                    else:
                        for item in list_item['items']:
                            para = doc.add_paragraph(item['text'], style='List Number')
            
            # Salva il documento
            output_path = os.path.join(output_directory, f"{base_name}_DSA_{timestamp}.docx")
            doc.save(output_path)
            
            logger.info(f"DOCX esportato: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Errore nell'export DOCX: {e}")
            raise
    
    def _setup_docx_styles(self, doc: Document, dsa_profile: DSAProfile):
        """Configura gli stili DOCX per il profilo DSA"""
        styles = doc.styles
        
        # Stile per il corpo del testo
        if 'DSA Body' not in [style.name for style in styles]:
            body_style = styles.add_style('DSA Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = dsa_profile.font
            body_font.size = Pt(dsa_profile.fontSize)
            body_para = body_style.paragraph_format
            body_para.line_spacing = dsa_profile.lineHeight
            body_para.space_after = Pt(dsa_profile.paragraphSpacing)
            body_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Stile per i titoli
        for i in range(1, 7):
            heading_style = styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = dsa_profile.font
            heading_font.size = Pt(dsa_profile.fontSize + (7 - i) * 2)
            heading_para = heading_style.paragraph_format
            heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading_para.space_after = Pt(dsa_profile.paragraphSpacing * 1.5)
    
    async def _export_pdf(
        self,
        structured_content: Dict[str, Any],
        dsa_profile: DSAProfile,
        output_directory: str,
        base_name: str,
        timestamp: str
    ) -> str:
        """Esporta in formato PDF (temporaneamente disabilitato)"""
        try:
            # Per ora, crea un file HTML che può essere convertito in PDF manualmente
            html_content = self._generate_html(structured_content, dsa_profile)
            css_content = self._generate_pdf_css(dsa_profile)
            
            # Salva come HTML
            output_path = os.path.join(output_directory, f"{base_name}_DSA_{timestamp}.html")
            
            full_html = f"""
<!DOCTYPE html>
<html lang="it">
<head>
    <meta charset="UTF-8">
    <title>{structured_content.get('title', 'Documento DSA')}</title>
    <style>{css_content}</style>
</head>
<body>
{html_content}
</body>
</html>
            """
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            logger.info(f"HTML esportato (per conversione PDF): {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Errore nell'export PDF: {e}")
            raise
    
    def _generate_html(self, structured_content: Dict[str, Any], dsa_profile: DSAProfile) -> str:
        """Genera HTML dal contenuto strutturato"""
        html_parts = ['<!DOCTYPE html>', '<html lang="it">', '<head>', '<meta charset="UTF-8">', '</head>', '<body>']
        
        # Aggiungi il titolo
        if structured_content.get('title'):
            html_parts.append(f'<h1>{structured_content["title"]}</h1>')
        
        # Aggiungi le sezioni
        for section in structured_content.get('sections', []):
            level = min(section['level'], 6)
            html_parts.append(f'<h{level}>{section["title"]}</h{level}>')
            
            for content_line in section['content']:
                if content_line.strip():
                    html_parts.append(f'<p>{content_line}</p>')
        
        # Aggiungi i paragrafi
        for paragraph in structured_content.get('paragraphs', []):
            if paragraph.strip():
                html_parts.append(f'<p>{paragraph}</p>')
        
        # Aggiungi gli elenchi
        for list_item in structured_content.get('lists', []):
            if list_item['items']:
                if list_item['type'] == 'bullet':
                    html_parts.append('<ul>')
                    for item in list_item['items']:
                        html_parts.append(f'<li>{item["text"]}</li>')
                    html_parts.append('</ul>')
                else:
                    html_parts.append('<ol>')
                    for item in list_item['items']:
                        html_parts.append(f'<li>{item["text"]}</li>')
                    html_parts.append('</ol>')
        
        html_parts.extend(['</body>', '</html>'])
        return '\n'.join(html_parts)
    
    def _generate_pdf_css(self, dsa_profile: DSAProfile) -> str:
        """Genera CSS per l'export PDF"""
        return f"""
        @page {{
            size: A4;
            margin: 2cm;
        }}
        
        body {{
            font-family: '{dsa_profile.font}', sans-serif;
            font-size: {dsa_profile.fontSize}px;
            line-height: {dsa_profile.lineHeight};
            color: {dsa_profile.textColor};
            background-color: {dsa_profile.backgroundColor};
            max-width: {dsa_profile.maxWidth}ch;
            margin: 0 auto;
            text-align: {dsa_profile.textAlign};
        }}
        
        h1, h2, h3, h4, h5, h6 {{
            font-family: '{dsa_profile.font}', sans-serif;
            color: {dsa_profile.textColor};
            margin-top: {dsa_profile.paragraphSpacing * 2}px;
            margin-bottom: {dsa_profile.paragraphSpacing}px;
        }}
        
        p {{
            margin-bottom: {dsa_profile.paragraphSpacing}px;
            text-align: {dsa_profile.textAlign};
        }}
        
        ul, ol {{
            margin-bottom: {dsa_profile.paragraphSpacing}px;
        }}
        
        li {{
            margin-bottom: {dsa_profile.paragraphSpacing / 2}px;
        }}
        
        a {{
            color: {dsa_profile.linkColor};
            text-decoration: underline;
        }}
        """
    
    async def _export_epub(
        self,
        structured_content: Dict[str, Any],
        dsa_profile: DSAProfile,
        output_directory: str,
        base_name: str,
        timestamp: str
    ) -> str:
        """Esporta in formato ePub"""
        try:
            # Crea un nuovo libro ePub
            book = epub.EpubBook()
            
            # Metadati del libro
            book.set_identifier(f"{base_name}_{timestamp}")
            book.set_title(structured_content.get('title', 'Documento DSA'))
            book.set_language('it')
            book.add_author('PDF DSA Converter')
            
            # Crea il contenuto HTML
            html_content = self._generate_html(structured_content, dsa_profile)
            
            # Crea il capitolo
            chapter = epub.EpubHtml(
                title='Contenuto',
                file_name='chapter.xhtml',
                lang='it'
            )
            chapter.content = html_content
            
            # Aggiungi il capitolo al libro
            book.add_item(chapter)
            
            # Crea la tabella dei contenuti
            book.toc = [chapter]
            
            # Aggiungi la navigazione
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Definisci lo spine
            book.spine = ['nav', chapter]
            
            # Salva il libro
            output_path = os.path.join(output_directory, f"{base_name}_DSA_{timestamp}.epub")
            epub.write_epub(output_path, book, {})
            
            logger.info(f"ePub esportato: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Errore nell'export ePub: {e}")
            raise
